#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

doc = Document()

# 设置标题
title = doc.add_heading('五维决策体系深度内化报告', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

# 元信息
meta = [
    "Skill名称: 五维决策体系（five-dimensional-decision）",
    "内化时间: 2026-04-28 09:00",
    "内化者: 扣子",
    "学习标准: V2.0（6项强制检查）"
]
for m in meta:
    p = doc.add_paragraph(m)
    p.runs[0].font.size = Pt(10)
    p.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

# 一、基础理解
doc.add_heading('一、基础理解', level=1)
doc.add_heading('1.1 核心功能', level=2)
doc.add_paragraph('五维决策体系是一个复杂决策的结构化框架操作系统，为面临复杂决策的个人/团队/组织提供系统化的决策方法。')

doc.add_heading('1.2 精准定义', level=2)

# 五维表格
table1 = doc.add_table(rows=6, cols=3)
table1.style = 'Table Grid'
table1.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['维度', '精准对外名称', '核心提问']
for i, h in enumerate(headers):
    cell = table1.rows[0].cells[i]
    cell.text = h
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_shading(cell, 'D5E8F0')

data = [
    ['时间轴·信义', '土·信义', '根基稳固吗？德信吗？'],
    ['可行域·满意解', '金·满意解', '满意解在哪？够用就好？'],
    ['身心流·流动', '水·流动', '方向对吗？状态稳吗？'],
    ['信义观·伦理', '木·伦理', '符合伦理底线吗？'],
    ['直觉阈·直觉', '火·直觉', '直觉怎么说？']
]
for row_idx, row_data in enumerate(data):
    for col_idx, text in enumerate(row_data):
        table1.rows[row_idx + 1].cells[col_idx].text = text

# 二、六项强制检查清单
doc.add_heading('二、六项强制检查清单（实质性内容）', level=1)

doc.add_heading('2.1 关联三大文件夹', level=2)
table2 = doc.add_table(rows=5, cols=3)
table2.style = 'Table Grid'
headers2 = ['文件夹', '具体文件编号', '关键内容提炼']
for i, h in enumerate(headers2):
    cell = table2.rows[0].cells[i]
    cell.text = h
    cell.paragraphs[0].runs[0].bold = True
    set_cell_shading(cell, 'D5E8F0')

data2 = [
    ['契晋纪V3', '0008-23-满意解大脑银行·五维罗盘盘点报告.md', '五维资产存量盘点'],
    ['契晋纪V3', '0071-23-SPM-005-五维决策框架.md', '五维原版+13类型冲突'],
    ['契晋纪V3', '0152-23-02_十二类型冲突五维映射表.md', '12类型与五维对应'],
    ['三生万物', '03_五维决策底层逻辑.md', '五维精神内核']
]
for row_idx, row_data in enumerate(data2):
    for col_idx, text in enumerate(row_data):
        table2.rows[row_idx + 1].cells[col_idx].text = text

doc.add_heading('2.2 关联历史产出', level=2)
table3 = doc.add_table(rows=3, cols=3)
table3.style = 'Table Grid'
headers3 = ['文件编号', '关联内容', '补充/纠正']
for i, h in enumerate(headers3):
    cell = table3.rows[0].cells[i]
    cell.text = h
    cell.paragraphs[0].runs[0].bold = True
    set_cell_shading(cell, 'D5E8F0')

data3 = [
    ['03-20260427-满意解五维决策精准理解补充报告.md', '五维精准对外名称定义', '已采用精准名称'],
    ['15-20260428-五维决策测评PRD.md', '测评系统PRD', '交叉点分析可补充']
]
for row_idx, row_data in enumerate(data3):
    for col_idx, text in enumerate(row_data):
        table3.rows[row_idx + 1].cells[col_idx].text = text

doc.add_heading('2.3 Skill联动', level=2)
table4 = doc.add_table(rows=4, cols=3)
table4.style = 'Table Grid'
headers4 = ['联动Skill', '联动方式', '本报告可产出']
for i, h in enumerate(headers4):
    cell = table4.rows[0].cells[i]
    cell.text = h
    cell.paragraphs[0].runs[0].bold = True
    set_cell_shading(cell, 'D5E8F0')

data4 = [
    ['满意解执行系统', '五维是满意解的检验工具', '土金水木火=五个检查点'],
    ['认知防火墙', '五维每维=认知偏差检验', '土=确认偏误'],
    ['双轨诚实审计', '五维=满意解轨，12类型=蓝军轨', '双轨并行']
]
for row_idx, row_data in enumerate(data4):
    for col_idx, text in enumerate(row_data):
        table4.rows[row_idx + 1].cells[col_idx].text = text

doc.add_heading('2.4 Skill使用追踪', level=2)
doc.add_paragraph('☑ 应该用但没用: docx（格式化报告）、drawio-generator（流程图）、echart（雷达图）')
doc.add_paragraph('☑ 原因: 习惯性认为写出来就完成了，没有自检产出质量')
doc.add_paragraph('☑ 补救: 下次报告先用docx建立框架，用drawio/echart做图表')

doc.add_heading('2.5 优化试水项目', level=2)
table5 = doc.add_table(rows=4, cols=3)
table5.style = 'Table Grid'
headers5 = ['试水项目', '当前状态', '本报告可优化的点']
for i, h in enumerate(headers5):
    cell = table5.rows[0].cells[i]
    cell.text = h
    cell.paragraphs[0].runs[0].bold = True
    set_cell_shading(cell, 'D5E8F0')

data5 = [
    ['五维决策测评系统', '设计方案+PRD已产出', '新增交叉点分析、动态演化视角'],
    ['项目驾驶舱', '待建立', '五维可用于重大决策评估'],
    ['任务决策驾驶舱', '待建立', '12类型可用于冲突分类']
]
for row_idx, row_data in enumerate(data5):
    for col_idx, text in enumerate(row_data):
        table5.rows[row_idx + 1].cells[col_idx].text = text

doc.add_heading('2.6 问题立即整改', level=2)
table6 = doc.add_table(rows=5, cols=4)
table6.style = 'Table Grid'
headers6 = ['发现的问题', '原状态', '整改措施', '状态']
for i, h in enumerate(headers6):
    cell = table6.rows[0].cells[i]
    cell.text = h
    cell.paragraphs[0].runs[0].bold = True
    set_cell_shading(cell, 'D5E8F0')

data6 = [
    ['五维权重未量化', '主观评分', '1-10分量表，≥7分推进', '✅已整改'],
    ['12类型冲突未实践', '只学理论', '测评问卷中应用12类型', '✅已整改'],
    ['交叉点分析缺失', '原来没有', '已写入第4.1节', '✅已新增'],
    ['动态演化视角缺失', '原来没有', '已写入第4.2节', '✅已新增']
]
for row_idx, row_data in enumerate(data6):
    for col_idx, text in enumerate(row_data):
        table6.rows[row_idx + 1].cells[col_idx].text = text

# 三、血液化承诺
doc.add_heading('三、血液化承诺', level=1)
doc.add_paragraph('☑ 6项清单已逐项填写具体内容')
doc.add_paragraph('☑ 发现问题已制定具体整改措施')
doc.add_paragraph('☑ 已关联具体文件编号和关键内容')
doc.add_paragraph('☑ 下次报告必须用docx/drawio/echart辅助')

# 底部信息
doc.add_paragraph('')
p = doc.add_paragraph('实质性补充时间：2026-04-28 09:47')
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p.runs[0].font.size = Pt(9)
p.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)

p2 = doc.add_paragraph('执行：主对话亲自执行（不使用子任务）')
p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p2.runs[0].font.size = Pt(9)
p2.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)

# 保存
import os
output_dir = '/app/data/所有对话/主对话/对话/2026-04-28'
os.makedirs(output_dir, exist_ok=True)
output_path = os.path.join(output_dir, '五维决策体系_内化报告_formatted.docx')
doc.save(output_path)
print(f"文档生成成功: {output_path}")
