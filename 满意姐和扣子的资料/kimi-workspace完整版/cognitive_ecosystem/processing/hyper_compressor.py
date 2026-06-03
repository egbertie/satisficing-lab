"""
三级蒸馏塔实现
L1: 结构拆解 -> L2: 语义向量化 -> L3: 认知晶体生成
带API故障时的本地启发式回退
"""

import os
import json
import asyncio
import re
import random
from pathlib import Path
from typing import List, Dict, Optional
from collections import deque
from docx import Document
import openai

from base.crystal_models import CognitiveCrystal, TotemType


class HyperCompressor:
    """
    超压缩引擎：将任意文档转化为认知晶体
    依赖：python-docx, openai（可选，失败时启发式回退）
    """

    def __init__(self,
                 embedding_model: str = "text-embedding-3-large",
                 llm_model: str = "gpt-4o",
                 compression_target: float = 0.15):
        self.embedding_model = embedding_model
        self.llm_model = llm_model
        self.compression_target = compression_target
        api_key = os.getenv("GITHUB_TOKEN")
        self.client = None
        if api_key:
            self.client = openai.AsyncOpenAI(
                base_url="https://models.inference.ai.azure.com",
                api_key=api_key
            )
        self.processing_queue = deque()

    async def process_docx(self, file_path: str, context_tags: List[str] = None) -> CognitiveCrystal:
        """主入口：DOCX -> 认知晶体"""
        raw_chunks = await self._l1_structure_extract(file_path)
        try:
            chunk_embeddings = await self._l2_semantic_vectorize(raw_chunks)
        except Exception as e:
            print(f"   ⚠️ Embedding API 失败 ({e})，启用本地回退...")
            chunk_embeddings = [{"embedding_id": f"emb-fallback-{i}"} for i in range(len(raw_chunks))]
        try:
            crystal = await self._l3_crystal_synthesis(
                source_path=file_path,
                chunks=raw_chunks,
                embeddings=chunk_embeddings,
                context_tags=context_tags or []
            )
        except Exception as e:
            print(f"   ⚠️ LLM API 失败 ({e})，启用启发式晶体合成...")
            crystal = self._l3_heuristic_synthesis(
                source_path=file_path,
                chunks=raw_chunks,
                embeddings=chunk_embeddings,
                context_tags=context_tags or []
            )
        return crystal

    async def _l1_structure_extract(self, file_path: str) -> List[Dict]:
        """L1: 结构拆解，保留语义块与表格"""
        doc = Document(file_path)
        chunks = []
        current_section = "无标题"
        for para in doc.paragraphs:
            style_name = para.style.name if para.style and para.style.name else "Normal"
            if style_name.startswith('Heading'):
                current_section = para.text
            if len(para.text.strip()) > 20:
                chunks.append({
                    "type": "paragraph",
                    "section": current_section,
                    "content": para.text,
                    "word_count": len(para.text)
                })
        for table_idx, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            chunks.append({
                "type": "table",
                "section": f"Table-{table_idx}",
                "content": json.dumps(table_data, ensure_ascii=False),
                "word_count": sum(len(str(cell)) for row in table_data for cell in row)
            })
        return chunks

    async def _l2_semantic_vectorize(self, chunks: List[Dict]) -> List[Dict]:
        """L2: 批量嵌入，使用GitHub Models"""
        if not chunks or not self.client:
            return []
        texts = [c["content"][:8000] for c in chunks]
        response = await self.client.embeddings.create(
            model=self.embedding_model,
            input=texts,
            dimensions=3072
        )
        for i, chunk in enumerate(chunks):
            chunk["embedding"] = response.data[i].embedding
            chunk["embedding_id"] = f"emb-{i}"
        return chunks

    async def _l3_crystal_synthesis(self,
                                    source_path: str,
                                    chunks: List[Dict],
                                    embeddings: List[Dict],
                                    context_tags: List[str]) -> CognitiveCrystal:
        """L3: GraphRAG风格认知晶体生成（LLM模式）"""
        if not self.client:
            raise RuntimeError("OpenAI client not available")
        synthesis_prompt = self._build_synthesis_prompt(chunks, context_tags)
        response = await self.client.chat.completions.create(
            model=self.llm_model,
            messages=[
                {"role": "system", "content": "你是一个认知蒸馏专家。将文档转化为结构化知识晶体，提取实体、关系、决策模式，并标记潜在矛盾。"},
                {"role": "user", "content": synthesis_prompt}
            ],
            response_format={"type": "json_object"},
            temperature=0.1
        )
        result = json.loads(response.choices[0].message.content)
        return self._build_crystal_from_result(source_path, chunks, embeddings, result, context_tags)

    def _l3_heuristic_synthesis(self,
                                 source_path: str,
                                 chunks: List[Dict],
                                 embeddings: List[Dict],
                                 context_tags: List[str]) -> CognitiveCrystal:
        """L3 启发式回退：基于规则提取晶体"""
        full_text = " ".join(c["content"] for c in chunks)[:5000]

        # 实体提取：找带引号的人名、书名号、大写组织名
        entities = list(set(
            re.findall(r'[《<\[]([^》>\]]+)[》>\]]', full_text) +
            re.findall(r'[\u4e00-\u9fa5]{2,4}?(?:教授|博士|先生|女士|所长|院长)', full_text)
        ))[:20]

        # 关系提取：简单 "A与B" / "A对B" 模式
        relations = []
        relation_patterns = [
            r'([\u4e00-\u9fa5]{2,8})与([\u4e00-\u9fa5]{2,8})',
            r'([\u4e00-\u9fa5]{2,8})对([\u4e00-\u9fa5]{2,8})'
        ]
        for pat in relation_patterns:
            for m in re.finditer(pat, full_text):
                relations.append({
                    "subject": m.group(1),
                    "predicate": "相关",
                    "object": m.group(2),
                    "confidence": 0.6
                })
        relations = relations[:10]

        # 决策模式：找 "应当"、"需要"、"必须" 句子
        decision_patterns = [
            s.strip() for s in re.findall(r'[^。]*(?:应当|需要|必须|应该|要)[^。]*。', full_text)
        ][:5]

        result = {
            "entities": entities,
            "relations": relations,
            "decision_patterns": decision_patterns,
            "contradictions": [],
            "triggers": context_tags,
            "abstract": full_text[:100]
        }
        return self._build_crystal_from_result(source_path, chunks, embeddings, result, context_tags)

    def _build_crystal_from_result(self, source_path, chunks, embeddings, result, context_tags):
        original_length = sum(c.get("word_count", 0) for c in chunks)
        compressed_length = len(json.dumps(result))
        ratio = compressed_length / original_length if original_length > 0 else 0
        totem_affinity = self._calculate_totem_affinity(result)
        return CognitiveCrystal(
            source_uris=[source_path],
            compression_ratio=min(1.0, ratio),
            primary_entities=result.get("entities", []),
            key_relations=result.get("relations", []),
            decision_patterns=result.get("decision_patterns", []),
            contradiction_flags=result.get("contradictions", []),
            activation_triggers=result.get("triggers", []) + context_tags,
            totem_affinity=totem_affinity,
            vector_fingerprint=embeddings[0]["embedding_id"] if embeddings else None
        )

    def _build_synthesis_prompt(self, chunks: List[Dict], context_tags: List[str]) -> str:
        """构建合成提示"""
        content_sample = "\n\n".join([
            f"[{c['section']}] {c['content'][:500]}"
            for c in chunks[:10]
        ])
        return f"""分析以下文档内容，提取结构化知识：

上下文标签：{', '.join(context_tags)}

文档内容：
{content_sample}

请输出JSON格式：
{{
    "entities": ["实体1", "实体2", ...],
    "relations": [
        {{"subject": "实体1", "predicate": "关系", "object": "实体2", "confidence": 0.9}}
    ],
    "decision_patterns": ["当X发生时，应采取Y", ...],
    "contradictions": ["与常识/其他文档冲突的点", ...],
    "triggers": ["激活此知识的场景关键词"],
    "abstract": "100字内核心摘要"
}}"""

    def _calculate_totem_affinity(self, result: Dict) -> Dict[str, float]:
        """基于内容计算与五路图腾的亲和度"""
        text = json.dumps(result, ensure_ascii=False)
        affinity = {
            TotemType.CONFUCIUS.value: 0.5,
            TotemType.SIMON.value: 0.5,
            TotemType.GUANYIN.value: 0.5,
            TotemType.LIUYUXI.value: 0.5,
            TotemType.HUINENG.value: 0.5
        }
        ethics_keywords = ["伦理", "道德", "长期", "信任", "责任", "隐私", "诚信"]
        if any(k in text for k in ethics_keywords):
            affinity[TotemType.CONFUCIUS.value] = 0.9
        risk_keywords = ["风险", "危机", "黑天鹅", "失败", "漏洞", "盲区"]
        if any(k in text for k in risk_keywords):
            affinity[TotemType.GUANYIN.value] = 0.9
        cost_keywords = ["优化", "成本", "预算", "效率", "ROI", "资源", "足够好"]
        if any(k in text for k in cost_keywords):
            affinity[TotemType.SIMON.value] = 0.9
        relation_keywords = ["关系", "网络", "社交", "合作", "伙伴", "连接"]
        if any(k in text for k in relation_keywords):
            affinity[TotemType.LIUYUXI.value] = 0.9
        innovation_keywords = ["创新", "突破", "范式", "重构", "新方案", "跃迁"]
        if any(k in text for k in innovation_keywords):
            affinity[TotemType.HUINENG.value] = 0.9
        return affinity

    async def batch_process(self,
                            directory: str,
                            pattern: str = "*.docx",
                            max_concurrent: int = 2) -> List[CognitiveCrystal]:
        """批量处理目录，带限流"""
        files = list(Path(directory).glob(pattern))
        semaphore = asyncio.Semaphore(max_concurrent)

        async def process_with_limit(f):
            async with semaphore:
                return await self.process_docx(str(f))

        tasks = [process_with_limit(f) for f in files]
        return await asyncio.gather(*tasks)
