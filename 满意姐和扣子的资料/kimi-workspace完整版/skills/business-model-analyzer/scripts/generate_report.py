#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
商业模式分析器 - Word报告生成器
基于SKILL框架生成商业模式调研报告
"""

import sys
import json
from datetime import datetime
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    print("[ERROR] 请先安装: pip install python-docx")
    sys.exit(1)


def create_cover(doc, company_name, date_str):
    """创建封面"""
    for _ in range(6):
        doc.add_paragraph()
    
    title_para = doc.add_paragraph()
    title_run = title_para.add_run("商业模式调研报告")
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 102, 204)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for _ in range(2):
        doc.add_paragraph()
    
    subtitle_para = doc.add_paragraph()
    subtitle_run = subtitle_para.add_run("基于SKILL框架的系统分析")
    subtitle_run.font.size = Pt(16)
    subtitle_run.font.color.rgb = RGBColor(102, 102, 102)
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    if company_name:
        company_para = doc.add_paragraph()
        company_run = company_para.add_run(f"分析对象：{company_name}")
        company_run.font.size = Pt(14)
        company_run.font.color.rgb = RGBColor(51, 51, 51)
        company_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for _ in range(4):
        doc.add_paragraph()
    
    date_para = doc.add_paragraph()
    date_run = date_para.add_run(date_str)
    date_run.font.size = Pt(12)
    date_run.font.color.rgb = RGBColor(128, 128, 128)
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()


def add_heading_custom(doc, text, level=1):
    """添加自定义标题"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        if level == 1:
            run.font.color.rgb = RGBColor(0, 102, 204)
            run.font.size = Pt(18)
        elif level == 2:
            run.font.color.rgb = RGBColor(0, 153, 204)
            run.font.size = Pt(14)
        else:
            run.font.color.rgb = RGBColor(51, 51, 51)
    return heading


def create_bmc_table(doc, bmc_data):
    """创建商业模式画布表格"""
    add_heading_custom(doc, "商业模式画布 (Business Model Canvas)", level=2)
    
    # 创建9宫格表格
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Light Grid Accent 1'
    
    # 填充画布内容
    blocks = [
        ("关键伙伴\nKP", bmc_data.get('kp', '')),
        ("关键业务\nKA", bmc_data.get('ka', '')),
        ("核心资源\nKR", bmc_data.get('kr', '')),
        ("价值主张\nVP", bmc_data.get('vp', '')),
        ("客户关系\nCR", bmc_data.get('cr', '')),
        ("渠道通路\nCH", bmc_data.get('ch', '')),
        ("客户细分\nCS", bmc_data.get('cs', '')),
        ("成本结构\nCS", bmc_data.get('cost', '')),
        ("收入来源\nRS", bmc_data.get('rs', '')),
    ]
    
    idx = 0
    for row in table.rows:
        for cell in row.cells:
            if idx < len(blocks):
                title, content = blocks[idx]
                cell.text = f"{title}\n\n{content}"
                idx += 1
    
    doc.add_paragraph()


def create_dimension_analysis(doc, dimension_data):
    """创建维度分析"""
    add_heading_custom(doc, "K - 研究维度分析", level=2)
    
    # 逻辑四维
    add_heading_custom(doc, "逻辑四维（硬实力）", level=3)
    logic_dimensions = dimension_data.get('logic', {})
    for key, value in logic_dimensions.items():
        para = doc.add_paragraph()
        label = para.add_run(f"{key}：")
        label.font.bold = True
        para.add_run(str(value))
        para.paragraph_format.line_spacing = 1.5
    
    doc.add_paragraph()
    
    # 深化三维
    add_heading_custom(doc, "深化三维（软实力）", level=3)
    deep_dimensions = dimension_data.get('deep', {})
    for key, value in deep_dimensions.items():
        para = doc.add_paragraph()
        label = para.add_run(f"{key}：")
        label.font.bold = True
        para.add_run(str(value))
        para.paragraph_format.line_spacing = 1.5
    
    doc.add_paragraph()


def create_position_analysis(doc, position_data):
    """创建产业链定位分析"""
    add_heading_custom(doc, "I - 产业链定位", level=2)
    
    position = position_data.get('position', '')
    para = doc.add_paragraph()
    label = para.add_run("产业链位置：")
    label.font.bold = True
    para.add_run(position)
    
    for key, value in position_data.items():
        if key != 'position':
            para = doc.add_paragraph()
            label = para.add_run(f"{key}：")
            label.font.bold = True
            para.add_run(str(value))
            para.paragraph_format.line_spacing = 1.5
    
    doc.add_paragraph()


def create_cultural_analysis(doc, cultural_data):
    """创建文化认知分析"""
    add_heading_custom(doc, "L - 文化主张与认知解决", level=2)
    
    # 4C框架
    add_heading_custom(doc, "4C价值主张", level=3)
    c4 = cultural_data.get('4c', {})
    for key, value in c4.items():
        para = doc.add_paragraph()
        label = para.add_run(f"{key}：")
        label.font.bold = True
        para.add_run(str(value))
        para.paragraph_format.line_spacing = 1.5
    
    doc.add_paragraph()
    
    # 认知路径
    add_heading_custom(doc, "认知路径", level=3)
    cognitive = cultural_data.get('cognitive', {})
    for key, value in cognitive.items():
        para = doc.add_paragraph()
        label = para.add_run(f"{key}：")
        label.font.bold = True
        para.add_run(str(value))
        para.paragraph_format.line_spacing = 1.5
    
    doc.add_paragraph()


def generate_report(data, output_path):
    """生成完整的商业模式调研报告"""
    doc = Document()
    
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    style.font.size = Pt(11)
    
    company_name = data.get('company_name', '')
    date_str = data.get('date', datetime.now().strftime('%Y年%m月%d日'))
    
    # 1. 封面
    create_cover(doc, company_name, date_str)
    
    # 2. 商业模式画布
    add_heading_custom(doc, "S - 结构板块分析", level=1)
    doc.add_paragraph()
    
    bmc = data.get('bmc', {})
    if bmc:
        create_bmc_table(doc, bmc)
    
    doc.add_page_break()
    
    # 3. 研究维度
    add_heading_custom(doc, "K - 研究维度分析", level=1)
    doc.add_paragraph()
    
    dimension = data.get('dimension', {})
    if dimension:
        create_dimension_analysis(doc, dimension)
    
    doc.add_page_break()
    
    # 4. 产业链定位
    add_heading_custom(doc, "I - 产业链定位", level=1)
    doc.add_paragraph()
    
    position = data.get('position', {})
    if position:
        create_position_analysis(doc, position)
    
    doc.add_page_break()
    
    # 5. 文化认知
    add_heading_custom(doc, "L - 文化主张与认知解决", level=1)
    doc.add_paragraph()
    
    cultural = data.get('cultural', {})
    if cultural:
        create_cultural_analysis(doc, cultural)
    
    doc.add_page_break()
    
    # 6. 综合评估
    add_heading_custom(doc, "综合评估与建议", level=1)
    doc.add_paragraph()
    
    evaluation = data.get('evaluation', {})
    if evaluation:
        for key, value in evaluation.items():
            para = doc.add_paragraph()
            label = para.add_run(f"{key}：")
            label.font.bold = True
            para.add_run(str(value))
            para.paragraph_format.line_spacing = 1.5
    
    doc.save(output_path)
    print(f"[OK] 商业模式调研报告已生成: {output_path}")


def main():
    if len(sys.argv) < 2:
        print("用法: python generate_report.py <json文件> [输出路径]")
        sys.exit(1)
    
    json_file = sys.argv[1]
    output_path = sys.argv[2] if len(sys.argv) > 2 else "商业模式调研报告.docx"
    
    with open(json_file, 'r', encoding='utf-8') as f:
        data = json.load(f)
    
    generate_report(data, output_path)


if __name__ == '__main__':
    main()
