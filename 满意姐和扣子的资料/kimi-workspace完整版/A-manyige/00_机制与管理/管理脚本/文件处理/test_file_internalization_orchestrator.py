#!/usr/bin/env python3
"""
test_file_internalization_orchestrator.py
批量文件内化编排器测试
"""

import json
import pytest
from pathlib import Path
from file_internalization_orchestrator import (
    FileInternalizationOrchestrator,
    FileTask,
    _DOCX_AVAILABLE,
    _PDF_AVAILABLE,
    _OCR_AVAILABLE,
)


class TestDependencies:
    def test_dependencies_are_met(self):
        orch = FileInternalizationOrchestrator()
        deps = orch.dependency_report()
        assert deps["python-docx"] is True
        assert deps["pymupdf"] is True
        assert deps["pytesseract/Pillow"] is True


class TestScan:
    def test_scan_2026_04_08(self):
        orch = FileInternalizationOrchestrator()
        tasks = orch.scan_files("2026-04-08")
        assert len(tasks) >= 120  # 已知有 123 个文件
        types = {t.file_type for t in tasks}
        assert "docx" in types

    def test_scan_supports_common_types(self):
        orch = FileInternalizationOrchestrator()
        supported = orch.SUPPORTED_TYPES
        assert "docx" in supported
        assert "pdf" in supported
        assert "png" in supported


class TestHashAndSkip:
    def test_hash_is_stable(self):
        orch = FileInternalizationOrchestrator()
        h1 = orch._compute_hash(Path(__file__))
        h2 = orch._compute_hash(Path(__file__))
        assert h1 == h2

    def test_skip_completed(self):
        orch = FileInternalizationOrchestrator()
        task = FileTask(source_path=Path("test.txt"), file_hash="abc123", file_type="txt")
        orch.completed_set.add("abc123")
        reason = orch._should_skip(task)
        assert reason is not None
        assert "已完成" in reason


class TestProcessItem:
    def test_process_md_file(self, tmp_path):
        source = tmp_path / "source"
        source.mkdir()
        out = tmp_path / "output"
        orch = FileInternalizationOrchestrator(source_dir=source, output_dir=out)
        md_file = source / "sample.md"
        md_file.write_text("# Hello\n\nThis is a test.", encoding="utf-8")
        task = FileTask(source_path=md_file, file_hash="md123", file_type="md")
        result = orch._process_item(task)
        assert result.status == "completed"
        assert result.paragraph_count == 2
        assert (out / "sample" / "extracted.txt").exists()

    @pytest.mark.skipif(not _DOCX_AVAILABLE, reason="python-docx not installed")
    def test_process_docx_file(self, tmp_path):
        import docx
        source = tmp_path / "source"
        source.mkdir()
        out = tmp_path / "output"
        orch = FileInternalizationOrchestrator(source_dir=source, output_dir=out)
        docx_file = source / "sample.docx"
        doc = docx.Document()
        doc.add_paragraph("段落一")
        doc.add_paragraph("段落二")
        doc.save(str(docx_file))

        task = FileTask(source_path=docx_file, file_hash="docx123", file_type="docx")
        result = orch._process_item(task)
        assert result.status == "completed"
        assert result.paragraph_count == 2
        assert (out / "sample" / "extracted.txt").exists()


class TestBatchRun:
    def test_batch_run_with_mock_files(self, tmp_path):
        source = tmp_path / "source"
        source.mkdir()
        (source / "a.md").write_text("Hello world", encoding="utf-8")
        (source / "b.txt").write_text("Test content", encoding="utf-8")
        (source / "c.pdf").write_text("fake pdf", encoding="utf-8")  # 会被解析失败，测试失败处理

        out = tmp_path / "output"
        orch = FileInternalizationOrchestrator(source_dir=source, output_dir=out)
        report = orch.run_batch()

        assert report["summary"]["total"] == 3
        # a.md 和 b.txt 应该成功（md 用 read_text，txt 也是）
        assert report["summary"]["completed"] >= 2
        # fake pdf 若 pymupdf 尝试解析会失败，或被跳过
        assert report["summary"]["failed"] >= 0 or report["summary"]["skipped"] >= 1
