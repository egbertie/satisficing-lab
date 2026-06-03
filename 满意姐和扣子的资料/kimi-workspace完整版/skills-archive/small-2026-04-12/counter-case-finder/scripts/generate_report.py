#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
反向案例发现器 - Word报告生成器
生成反向案例分析报告
"""

import sys
import json
from datetime import datetime
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    print("[ERROR] 请先安装: pip install python-docx")
    sys.exit(1)


def set_run_font(run, font_size=11, bold=False, color=(0, 0, 0)):
    """设置字体样式"""
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(*color)


def add_heading_custom(doc, text, level=1):
    """添加自定义标题"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        if level == 1:
            run.font.color.rgb = RGBColor(220, 20, 60)  # 深红色
            run.font.size = Pt(18)
        elif level == 2:
            run.font.color.rgb = RGBColor(178, 34, 34)  # 火砖色
            run.font.size = Pt(14)
        else:
            run.font.color.rgb = RGBColor(51, 51, 51)
    return heading


def add_info_box(doc, title, content, box_type="info"):
    """添加信息框"""
    colors = {
        "info": (70, 130, 180),      # 钢蓝色
        "warning": (255, 140, 0),    # 深橙色
        "danger": (220, 20, 60),     # 深红色
        "success": (34, 139, 34),    # 森林绿
    }
    color = colors.get(box_type, (70, 130, 180))
    
    para = doc.add_paragraph()
    title_run = para.add_run(f"【{title}】\n")
    set_run_font(title_run, font_size=12, bold=True, color=color)
    
    content_run = para.add_run(content)
    set_run_font(content_run, font_size=11)
    
    para.paragraph_format.line_spacing = 1.5
    para.paragraph_format.space_after = Pt(12)


def create_cover(doc, subject, date_str):
    """创建封面"""
    for _ in range(6):
        doc.add_paragraph()
    
    title_para = doc.add_paragraph()
    title_run = title_para.add_run("反向案例分析报告")
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(220, 20, 60)  # 深红色
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for _ in range(2):
        doc.add_paragraph()
    
    subtitle_para = doc.add_paragraph()
    subtitle_run = subtitle_para.add_run("从反面看正面，从失败学成功")
    subtitle_run.font.size = Pt(16)
    subtitle_run.font.color.rgb = RGBColor(102, 102, 102)
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    if subject:
        subject_para = doc.add_paragraph()
        subject_run = subject_para.add_run(f"分析主题：{subject}")
        subject_run.font.size = Pt(14)
        subject_run.font.color.rgb = RGBColor(51, 51, 51)
        subject_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for _ in range(4):
        doc.add_paragraph()
    
    date_para = doc.add_paragraph()
    date_run = date_para.add_run(date_str)
    date_run.font.size = Pt(12)
    date_run.font.color.rgb = RGBColor(128, 128, 128)
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()


def create_step1_section(doc, data):
    """创建第一步：定义正面"""
    add_heading_custom(doc, "第一步：定义正面", level=1)
    doc.add_paragraph()
    
    # 事件要点
    add_heading_custom(doc, "事件要点梳理", level=2)
    
    event_desc = data.get('event_description', '')
    if event_desc:
        para = doc.add_paragraph()
        label = para.add_run("事件描述：")
        label.font.bold = True
        para.add_run(event_desc)
        para.paragraph_format.line_spacing = 1.5
    
    # 核心目标
    goals = data.get('core_goals', [])
    if goals:
        para = doc.add_paragraph()
        label = para.add_run("核心目标：")
        label.font.bold = True
        for goal in goals:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(goal)
    
    # 关键成功因素
    factors = data.get('success_factors', [])
    if factors:
        para = doc.add_paragraph()
        label = para.add_run("关键成功因素：")
        label.font.bold = True
        for factor in factors:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(factor)
    
    doc.add_paragraph()
    
    # 多维定义
    add_heading_custom(doc, "\"正面\"的多维定义", level=2)
    
    dimensions = [
        ("结果维度", data.get('dimension_result', '')),
        ("过程维度", data.get('dimension_process', '')),
        ("条件维度", data.get('dimension_condition', '')),
        ("时间维度", data.get('dimension_time', '')),
    ]
    
    for name, content in dimensions:
        if content:
            add_heading_custom(doc, name, level=3)
            para = doc.add_paragraph(content)
            para.paragraph_format.line_spacing = 1.5
    
    # 正面定义总结
    summary = data.get('positive_definition_summary', '')
    if summary:
        add_info_box(doc, "正面定义总结", summary, box_type="success")
    
    doc.add_page_break()


def create_step2_section(doc, data):
    """创建第二步：反面案例"""
    add_heading_custom(doc, "第二步：反面案例", level=1)
    doc.add_paragraph()
    
    cases = data.get('counter_cases', [])
    for i, case in enumerate(cases, 1):
        add_heading_custom(doc, f"反面案例{i}：{case.get('name', '未命名')}", level=2)
        
        # 基本信息
        add_heading_custom(doc, "基本信息", level=3)
        basic_info = case.get('basic_info', {})
        for key, value in basic_info.items():
            para = doc.add_paragraph()
            label = para.add_run(f"{key}：")
            label.font.bold = True
            para.add_run(str(value))
        
        # 背景情境
        background = case.get('background', '')
        if background:
            add_heading_custom(doc, "背景情境", level=3)
            para = doc.add_paragraph(background)
            para.paragraph_format.line_spacing = 1.5
        
        # 决策过程
        decision = case.get('decision', {})
        if decision:
            add_heading_custom(doc, "决策过程", level=3)
            for key, value in decision.items():
                para = doc.add_paragraph()
                label = para.add_run(f"{key}：")
                label.font.bold = True
                para.add_run(str(value))
        
        # 结果呈现
        results = case.get('results', {})
        if results:
            add_heading_custom(doc, "结果呈现", level=3)
            for key, value in results.items():
                para = doc.add_paragraph()
                label = para.add_run(f"{key}：")
                label.font.bold = True
                para.add_run(str(value))
        
        # 与正面定义的对比
        comparison = case.get('comparison', '')
        if comparison:
            add_info_box(doc, "与正面定义的对比", comparison, box_type="warning")
        
        doc.add_paragraph()
    
    doc.add_page_break()


def create_step3_section(doc, data):
    """创建第三步：总结建议"""
    add_heading_custom(doc, "第三步：总结建议", level=1)
    doc.add_paragraph()
    
    # 失败原因归类
    add_heading_custom(doc, "失败原因归类", level=2)
    
    failure_reasons = data.get('failure_reasons', {})
    for category, reasons in failure_reasons.items():
        add_heading_custom(doc, category, level=3)
        for reason in reasons:
            para = doc.add_paragraph(style='List Bullet')
            para.add_run(reason)
    
    doc.add_paragraph()
    
    # 核心建议
    add_heading_custom(doc, "核心建议", level=2)
    
    # 红线
    red_lines = data.get('red_lines', [])
    if red_lines:
        add_heading_custom(doc, "🚫 必须避免的红线", level=3)
        for line in red_lines:
            para = doc.add_paragraph(style='List Bullet')
            para.add_run(line)
        doc.add_paragraph()
    
    # 条件
    conditions = data.get('must_conditions', [])
    if conditions:
        add_heading_custom(doc, "✅ 必须满足的条件", level=3)
        for condition in conditions:
            para = doc.add_paragraph(style='List Bullet')
            para.add_run(condition)
        doc.add_paragraph()
    
    # 策略
    strategies = data.get('strategies', [])
    if strategies:
        add_heading_custom(doc, "🎯 建议采取的策略", level=3)
        for strategy in strategies:
            para = doc.add_paragraph(style='List Bullet')
            para.add_run(strategy)
        doc.add_paragraph()
    
    # 预警指标
    warnings = data.get('warning_indicators', [])
    if warnings:
        add_heading_custom(doc, "⚠️ 风险预警指标", level=3)
        for warning in warnings:
            para = doc.add_paragraph(style='List Bullet')
            para.add_run(warning)
        doc.add_paragraph()
    
    # 行动清单
    add_heading_custom(doc, "行动清单", level=2)
    
    actions = data.get('action_items', {})
    for period, items in actions.items():
        add_heading_custom(doc, period, level=3)
        for item in items:
            para = doc.add_paragraph(style='List Bullet')
            para.add_run(item)
    
    doc.add_page_break()


def generate_report(data, output_path):
    """生成完整的反向案例分析报告"""
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    style.font.size = Pt(11)
    
    # 提取数据
    subject = data.get('subject', '')
    date_str = data.get('date', datetime.now().strftime('%Y年%m月%d日'))
    
    # 1. 封面
    create_cover(doc, subject, date_str)
    
    # 2. 执行摘要
    add_heading_custom(doc, "执行摘要", level=1)
    
    # 核心问题
    core_problem = data.get('core_problem', '')
    if core_problem:
        add_info_box(doc, "用户核心问题", core_problem, box_type="info")
    
    # 关键发现
    key_findings = data.get('key_findings', [])
    if key_findings:
        add_heading_custom(doc, "关键发现", level=2)
        for finding in key_findings:
            para = doc.add_paragraph(style='List Bullet')
            para.add_run(finding)
        doc.add_paragraph()
    
    # 核心建议
    core_recommendations = data.get('core_recommendations', [])
    if core_recommendations:
        add_heading_custom(doc, "核心建议", level=2)
        for rec in core_recommendations:
            para = doc.add_paragraph(style='List Bullet')
            para.add_run(rec)
    
    doc.add_page_break()
    
    # 3. 第一步
    create_step1_section(doc, data)
    
    # 4. 第二步
    create_step2_section(doc, data)
    
    # 5. 第三步
    create_step3_section(doc, data)
    
    # 保存文档
    doc.save(output_path)
    print(f"[OK] 反向案例分析报告已生成: {output_path}")


def main():
    """主函数"""
    if len(sys.argv) < 2:
        print("用法: python generate_report.py <json文件> [输出路径]")
        sys.exit(1)
    
    json_file = sys.argv[1]
    output_path = sys.argv[2] if len(sys.argv) > 2 else "反向案例分析报告.docx"
    
    # 读取JSON数据
    with open(json_file, 'r', encoding='utf-8') as f:
        data = json.load(f)
    
    # 生成报告
    generate_report(data, output_path)


if __name__ == '__main__':
    main()
